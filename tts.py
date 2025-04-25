import speech_recognition as sr
from docx import Document

def listen_and_transcribe():
    recognizer = sr.Recognizer()
    mic = sr.Microphone()
    doc = Document()
    doc.add_heading('Voice Transcription', 0)

    print("🎤 Speak clearly. Say 'stop recording' to finish.\n")

    with mic as source:
        recognizer.adjust_for_ambient_noise(source)

        while True:
            print("Listening...")
            audio = recognizer.listen(source)

            try:
                text = recognizer.recognize_google(audio)
                print(f"You said: {text}")

                if "stop recording" in text.lower():
                    print("🛑 Stopping transcription.")
                    break

                doc.add_paragraph(text)

            except sr.UnknownValueError:
                print("❌ Sorry, couldn't understand. Try again.")
            except sr.RequestError as e:
                print(f"⚠️ Could not request results; {e}")

    doc.save("Voice_Notes.docx")
    print("✅ Saved as Voice_Notes.docx")

if __name__ == "__main__":
    listen_and_transcribe()
