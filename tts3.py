import speech_recognition as sr
from docx import Document

def process_text(text):
    # Replace spoken punctuation with actual punctuation
    replacements = {
        " comma": ",",
        " period": ".",
        " exclamation mark": "!",
        " question mark": "?",
    }

    for word, symbol in replacements.items():
        text = text.replace(word, symbol)

    return text

def listen_and_transcribe():
    recognizer = sr.Recognizer()
    mic = sr.Microphone()
    doc = Document()

    print("🎤 Speak clearly. Say 'stop recording' to finish.")
    print("💡 Use commands: 'title [text]', 'new line', 'comma', 'period', etc.\n")

    with mic as source:
        recognizer.adjust_for_ambient_noise(source)

        while True:
            print("Listening...")
            audio = recognizer.listen(source)

            try:
                text = recognizer.recognize_google(audio).lower()
                print(f"You said: {text}")

                if "stop recording" in text:
                    print("🛑 Stopping transcription.")
                    break

                # Insert title
                if text.startswith("title "):
                    title_text = text[6:]
                    title_text = process_text(title_text)
                    doc.add_heading(title_text, level=1)
                    print(f"📝 Title added: {title_text}")

                # Insert new paragraph
                elif text == "new line" or text == "newline":
                    doc.add_paragraph("")  # Adds a blank paragraph for spacing
                    print("↩️ New line added.")

                else:
                    processed_text = process_text(text)
                    doc.add_paragraph(processed_text)
                    print(f"📄 Text added: {processed_text}")

            except sr.UnknownValueError:
                print("❌ Couldn’t understand. Try again.")
            except sr.RequestError as e:
                print(f"⚠️ Could not request results; {e}")

    doc.save("Voice_Notes_Formatted.docx")
    print("✅ Saved as Voice_Notes_Formatted.docx")

if __name__ == "__main__":
    listen_and_transcribe()
