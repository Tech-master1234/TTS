import speech_recognition as sr
from docx import Document

def listen_and_transcribe():
    recognizer = sr.Recognizer()
    mic = sr.Microphone()
    doc = Document()

    print("🎤 Speak clearly. Say 'stop recording' to finish.\n")
    print("💡 Say 'title [your title]' to insert a heading.\n")

    with mic as source:
        recognizer.adjust_for_ambient_noise(source)

        while True:
            print("Listening...")
            audio = recognizer.listen(source)

            try:
                text = recognizer.recognize_google(audio)
                print(f"You said: {text}")

                if "stop recording" in text.lower():
                    print("🛑 Stopping transcription.")
                    break

                # Handle title insertion
                if text.lower().startswith("title "):
                    title_text = text[6:]  # Everything after the word "title"
                    doc.add_heading(title_text, level=1)
                    print(f"📝 Title added: {title_text}")
                else:
                    doc.add_paragraph(text)

            except sr.UnknownValueError:
                print("❌ Sorry, couldn't understand. Try again.")
            except sr.RequestError as e:
                print(f"⚠️ Could not request results; {e}")

    doc.save("Voice_Notes_with_Titles.docx")
    print("✅ Saved as Voice_Notes_with_Titles.docx")

if __name__ == "__main__":
    listen_and_transcribe()
